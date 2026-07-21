from pathlib import Path

from docx import Document

from app.importers.docx_parser import parse_docx


FIELDS = (
    ("基本信息", "类型：动漫名 ｜ 难度：★★☆☆☆ ｜ 答案字数：3"),
    ("图1生图提示词", "题{code}图1"),
    ("图1题面提示", "这是图1"),
    ("图2生图提示词", "题{code}图2"),
    ("图2题面填空", "这是 ＿ ＿ ＿"),
    ("答案", "第{answer}题"),
    ("数字声调拼音", "di4 ti2"),
    ("谜底拆解", "测试拆解"),
)


def add_question(document: Document, code: str, answer: str, *, unknown: bool = False) -> None:
    document.add_heading(f"{code}｜第{answer}题", level=2)
    table = document.add_table(rows=0, cols=2)
    for label, template in FIELDS:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = template.format(code=code, answer=answer)
    if unknown:
        row = table.add_row().cells
        row[0].text = "未来字段"
        row[1].text = "应保留为警告"
    document.add_paragraph("制作检查：不要串题。")


def build_fixture(path: Path) -> None:
    document = Document()
    document.add_heading("说明", level=1)
    intro = document.add_table(rows=1, cols=2)
    intro.cell(0, 0).text = "版本"
    intro.cell(0, 1).text = "测试"
    add_question(document, "001", "一", unknown=True)
    document.add_page_break()
    document.add_paragraph("")
    add_question(document, "002", "二")
    document.add_paragraph("盲测记录")
    trailing = document.add_table(rows=1, cols=2)
    trailing.cell(0, 0).text = "题号"
    trailing.cell(0, 1).text = "结果"
    document.save(path)


def test_docx_parser_keeps_fields_with_their_question(tmp_path: Path) -> None:
    path = tmp_path / "题库.docx"
    build_fixture(path)

    preview = parse_docx(path)

    assert [question.code for question in preview.questions] == ["001", "002"]
    assert preview.questions[0].image1_prompt == "题001图1"
    assert preview.questions[0].image2_prompt == "题001图2"
    assert preview.questions[1].image1_prompt == "题002图1"
    assert preview.questions[1].answer == "第二题"
    assert preview.questions[0].quality_check == "制作检查：不要串题。"
    assert any(issue.code == "UNKNOWN_FIELD" for issue in preview.issues)
    assert not any(issue.field == "题号" for issue in preview.issues)

